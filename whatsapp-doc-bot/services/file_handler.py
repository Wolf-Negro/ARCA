import io
import logging
from typing import Optional
from services.openai_service import read_image_content

logger = logging.getLogger(__name__)

SUPPORTED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp", "image/bmp"}
SUPPORTED_DOC_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/plain",
    "text/csv",
}


def process_incoming_file(
    file_bytes: bytes,
    mime_type: str,
    filename: str = "",
) -> str:
    """
    Extrae texto del archivo según su tipo MIME.
    Retorna el texto extraído o una descripción del contenido.
    """
    mime_type = mime_type.lower().split(";")[0].strip()

    if mime_type == "application/pdf":
        return _extract_pdf(file_bytes)

    if mime_type in SUPPORTED_IMAGE_TYPES:
        return read_image_content(file_bytes, mime_type)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_xlsx(file_bytes)

    if mime_type in ("text/plain", "text/csv"):
        return _extract_text(file_bytes)

    logger.warning("Tipo MIME no soportado directamente: %s (%s)", mime_type, filename)
    return f"Archivo: {filename}. Tipo: {mime_type}. Contenido no extraíble automáticamente."


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for i, page in enumerate(reader.pages):
            if i >= 20:  # límite de páginas para no exceder tokens
                break
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())

        content = "\n\n".join(pages_text)
        if not content.strip():
            return "PDF sin texto extraíble (posiblemente escaneado)"
        logger.info("PDF extraído: %d caracteres de %d páginas", len(content), len(reader.pages))
        return content[:10000]
    except Exception as e:
        logger.error("Error extrayendo PDF: %s", e)
        return "No se pudo extraer el contenido del PDF"


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # También extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = "\n".join(paragraphs)
        logger.info("DOCX extraído: %d caracteres", len(content))
        return content[:10000] if content else "Documento Word vacío"
    except Exception as e:
        logger.error("Error extrayendo DOCX: %s", e)
        return "No se pudo extraer el contenido del documento Word"


def _extract_xlsx(file_bytes: bytes) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        rows_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text.append(f"[Hoja: {sheet_name}]")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= 100:  # límite de filas
                    break
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    rows_text.append(" | ".join(cells))
                    row_count += 1

        content = "\n".join(rows_text)
        logger.info("XLSX extraído: %d caracteres", len(content))
        return content[:10000] if content else "Archivo Excel vacío"
    except Exception as e:
        logger.error("Error extrayendo XLSX: %s", e)
        return "No se pudo extraer el contenido del archivo Excel"


def _extract_text(file_bytes: bytes) -> str:
    try:
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                content = file_bytes.decode(encoding)
                logger.info("Texto extraído: %d caracteres (encoding: %s)", len(content), encoding)
                return content[:10000]
            except UnicodeDecodeError:
                continue
        return "Archivo de texto con encoding no compatible"
    except Exception as e:
        logger.error("Error extrayendo texto: %s", e)
        return "No se pudo leer el archivo de texto"


def get_file_extension(mime_type: str) -> str:
    extensions = {
        "application/pdf": ".pdf",
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/gif": ".gif",
        "image/webp": ".webp",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "application/vnd.ms-excel": ".xls",
        "text/plain": ".txt",
        "text/csv": ".csv",
    }
    mime_clean = mime_type.lower().split(";")[0].strip()
    return extensions.get(mime_clean, "")


def is_supported_file(mime_type: str) -> bool:
    mime_clean = mime_type.lower().split(";")[0].strip()
    return mime_clean in SUPPORTED_IMAGE_TYPES | SUPPORTED_DOC_TYPES
